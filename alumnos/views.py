import datetime
import os
from io import BytesIO
from copy import deepcopy

from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.conf import settings
from .forms import AlumnoForm
from .models import Alumno

from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

MESES = ['enero','febrero','marzo','abril','mayo','junio',
         'julio','agosto','septiembre','octubre','noviembre','diciembre']

TEMPLATES_DOCX = os.path.join(settings.BASE_DIR, 'templates_docx')


# ──────────────────────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────────────────────

def _today_str():
    hoy = datetime.date.today()
    return f"{hoy.day} de {MESES[hoy.month - 1]} de {hoy.year}"


def _set_run_text(run, text):
    run.text = text


def _safe(value, default=''):
    return (value or default).strip()


# ──────────────────────────────────────────────────────────────
# VISTAS PRINCIPALES
# ──────────────────────────────────────────────────────────────

def seleccion(request):
    return render(request, 'alumnos/seleccion.html')


def registro(request, facultad='MEDICINA'):
    facultad = facultad.upper()
    if request.method == 'POST':
        form = AlumnoForm(request.POST, request.FILES)
        if form.is_valid():
            alumno = form.save(commit=False)
            alumno.facultad = facultad
            alumno.save()
            return redirect('gracias', alumno_id=alumno.id)
    else:
        form = AlumnoForm()
    facultad_display = 'Facultad de Medicina' if facultad == 'MEDICINA' else 'Facultad de Química'
    return render(request, 'alumnos/registro.html', {
        'form': form,
        'facultad': facultad,
        'facultad_display': facultad_display,
    })


def gracias(request, alumno_id):
    alumno = Alumno.objects.get(id=alumno_id)
    return render(request, 'alumnos/gracias.html', {'alumno': alumno})


# ──────────────────────────────────────────────────────────────
# CONSTANCIA NO ADEUDO  →  constanciaedit.docx
# ──────────────────────────────────────────────────────────────

def generar_pdf(request, alumno_id):
    alumno = Alumno.objects.get(id=alumno_id)
    doc = Document(os.path.join(TEMPLATES_DOCX, 'constanciaedit.docx'))

    nombre  = _safe(alumno.nombre).upper()
    carrera = alumno.get_carrera_display().upper() if alumno.carrera else ''
    cuenta  = _safe(alumno.numero_cuenta)

    # Párrafo 9: "...------(Nombre)---------, pasante de la ----(LICENCIATURA)---,
    #             con número de cuenta---(NUMERO DE CUENTA)-------..."
    para = doc.paragraphs[9]
    runs = para.runs
    # Runs: [2]="  ------", [3]="(Nombre)---------", [5]=" ----(LICENCIATURA)---",
    #        [7]="---(NUMERO DE ", [8]="CUENTA)-------"
    if len(runs) > 8:
        runs[2].text = ' '
        runs[3].text = nombre
        runs[5].text = ' ' + carrera
        runs[7].text = ''
        runs[8].text = cuenta

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="constancia_{cuenta}.docx"'
    return response


# ──────────────────────────────────────────────────────────────
# REGISTRO MATERIAL  →  regristromaterial.docx
# ──────────────────────────────────────────────────────────────

def generar_pdf2(request, alumno_id):
    alumno = Alumno.objects.get(id=alumno_id)
    doc = Document(os.path.join(TEMPLATES_DOCX, 'regristromaterial.docx'))

    nombre = _safe(alumno.nombre).upper()
    tema   = _safe(alumno.tema)

    # Párrafo [0]: nombre del alumno
    para0 = doc.paragraphs[0]
    if para0.runs:
        para0.runs[0].text = nombre
    else:
        para0.add_run(nombre)

    # Tabla [0], fila 1 (primera fila de datos): Título | Autor | Edición | Editorial
    if doc.tables:
        table = doc.tables[0]
        if len(table.rows) > 1:
            row = table.rows[1]
            def _write_cell(cell, text):
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].text = text
                else:
                    cell.paragraphs[0].add_run(text)

            _write_cell(row.cells[0], tema)
            _write_cell(row.cells[1], nombre)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="registro_material_{alumno.numero_cuenta}.docx"'
    return response


# ──────────────────────────────────────────────────────────────
# CARTA DE AUTORIZACIÓN  →  Word generado
# ──────────────────────────────────────────────────────────────

def _add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11 if level == 1 else 10)
    return p


def _add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Courier New'
    return p


def _add_field(doc, label, value, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(label + ' ')
    r1.font.size = Pt(size)
    r1.font.name = 'Courier New'
    r2 = p.add_run(value if value else '_' * 30)
    r2.bold = bool(value)
    r2.font.size = Pt(size)
    r2.font.name = 'Courier New'
    return p


def generar_carta(request, alumno_id):
    alumno = Alumno.objects.get(id=alumno_id)
    fecha   = _today_str()
    nombre  = _safe(alumno.nombre)
    cuenta  = _safe(alumno.numero_cuenta)
    tema    = _safe(alumno.tema)
    carrera = alumno.get_carrera_display() if alumno.carrera else ''
    tel     = _safe(alumno.telefono)
    correo  = _safe(alumno.correo)
    domicilio = _safe(alumno.domicilio)
    director  = _safe(alumno.director)

    doc = Document()
    # Márgenes
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    def h(text, center=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Courier New'
        return p

    def body(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = 'Courier New'
        return p

    def right(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = 'Courier New'
        return p

    def sp():
        doc.add_paragraph()

    # ── Encabezado ──
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hdr.add_run('Universidad Autónoma del Estado de México')
    r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph().add_run('UAEM').font.size = Pt(9)

    # ── PÁGINA 1 ──
    right(f'Toluca, México; a {fecha}')
    sp()
    _add_heading(doc, 'Carta de autorización para la incorporación de\n'
                       'objetos digitales en el Repositorio Institucional de\n'
                       'la Universidad Autónoma del Estado de México.')
    sp()
    h('UNIVERSIDAD AUTÓNOMA DEL ESTADO DE MÉXICO')
    h('P R E S E N T E')
    sp()
    body(
        f'El/la/los que suscribe/n {nombre}, con '
        'fundamento en los artículos 13 fracción I, 18, 21 22, 27, 30 y demás '
        'aplicables de la Ley Federal del Derecho de Autor y su Reglamento vigentes, '
        'firmo/mamos la presente Licencia de Uso Gratuita, No Exclusiva y No '
        'remunerada para la incorporación al Repositorio Institucional de la '
        'Universidad Autónoma del Estado de México de la obra literaria (artículo, '
        'capítulo de libro, libro, tesis de posgrado, entre otros.) que lleva por '
        f'título {tema}'
    )
    sp()
    body(
        'Asimismo, declaro/ramos bajo protesta de decir verdad ser el/la/los '
        'autor/a/res y/o legítimo/a/s titular/es de la obra literaria y sus '
        'derivados visuales; y que responderé/remos de la autoría/titularidad, '
        'originalidad y nivel de acceso de la obra de mérito y del ejercicio '
        'pacífico de los derechos que se licencian en este acto, manifestando que '
        'no existe ninguna otra persona física o moral a la que le pertenezcan; '
        'por lo cual libero/ramos en este acto de toda responsabilidad a la '
        'Universidad Autónoma del Estado de México, así como de cualquier demanda '
        'o reclamación que llegara a formular alguna persona física o moral que '
        'considere vulnerados sus derechos o que se suponga con derecho sobre la '
        'obra mencionada, asumiendo todas las consecuencias legales y económicas '
        'a que hubiera lugar.'
    )
    sp()
    body(
        'Por lo anterior, autorizo que la Oficina de Conocimiento Abierto '
        'perteneciente a esta Máxima Casa de Estudios, realice lo propio para '
        'el almacenamiento, preservación y difusión de la obra, con fines '
        'académicos y culturales en formato de acceso abierto y sin fines de lucro '
        'en los términos siguientes:'
    )
    sp()
    h('1. De los Derechos de Autor.')
    sp()
    body(
        'Reconozco la importancia de protección de mi obra y el movimiento de '
        'Acceso Abierto del cual forma parte la Universidad Autónoma del Estado '
        'de México, por lo tanto conozco y acepto que mi obra esté protegida '
        'bajo una de las Licencia Creative Commons que a continuación se listan, '
        'marcando con una "X" del lado izquierdo la que será aplicable a mi obra:'
    )

    # ── PÁGINA 2 – Licencias CC ──
    doc.add_page_break()
    cc_items = [
        ('Reconocimiento (BY)',
         'El autor permite copiar, reproducir, distribuir, comunicar públicamente la obra, '
         'realizar obras derivadas y hacer uso comercial, siempre citando al autor original.'),
        ('Reconocimiento - Sin obra derivada (BY-ND)',
         'Permite uso comercial citando al autor. No permite obra derivada.'),
        ('Reconocimiento - No comercial - Sin obra derivada (BY-NC-ND)',
         'Permite difusión citando al autor. No permite obra derivada ni uso comercial.'),
        ('Reconocimiento - No comercial (BY-NC)',
         'Permite obra derivada citando al autor. No permite uso comercial.'),
        ('Reconocimiento - No comercial - Compartir igual (BY-NC-SA)',
         'Permite obra derivada bajo misma licencia, citando al autor. No permite uso comercial.'),
        ('Reconocimiento - Compartir igual (BY-SA)',
         'Permite uso comercial y obra derivada con misma licencia, citando al autor.'),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr_row = tbl.rows[0].cells
    hdr_row[0].text = ''
    hdr_row[1].text = 'Licencia'
    hdr_row[2].text = 'Icono'
    for name, desc in cc_items:
        row = tbl.add_row().cells
        row[0].text = '□'
        p = row[1].paragraphs[0]
        p.add_run(name + ': ').bold = True
        p.add_run(desc)
        row[2].text = 'CC ' + name.split('(')[-1].replace(')', '')
    sp()
    h('2. De la Difusión del producto')
    sp()
    body(
        'El nivel de acceso en mi obra definirá la parcialidad o totalidad de '
        'acceso a los datos y documento a texto completo para su visibilidad '
        'en el Repositorio Institucional, por lo que la aplicable a mi obra, '
        'es el señalada del lado izquierdo en esta sección:'
    )

    # ── PÁGINA 3 – Nivel de acceso ──
    doc.add_page_break()
    access_items = [
        ('a. Abierto',
         'Permite que los metadatos y el documento a texto completo sean visualizados '
         'y descargados por cualquier usuario de manera libre y sin costo.'),
        ('b. Restringido',
         'El documento no se muestra al público. Los metadatos son visibles '
         'a petición del depositante. Se notifica al autor si alguien solicita acceso.'),
        ('c. Embargado',
         'Oculta el documento por un periodo definido. Al vencer el embargo '
         'el acceso cambia automáticamente a "acceso abierto".'),
        ('d. Cerrado',
         'El depósito no aparece en búsquedas. El documento y los metadatos '
         'NO son visibles para los usuarios.'),
    ]
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = 'Table Grid'
    hdr2 = tbl2.rows[0].cells
    hdr2[0].text = ''
    hdr2[1].text = 'Nivel de acceso'
    for name, desc in access_items:
        row = tbl2.add_row().cells
        row[0].text = '□'
        p = row[1].paragraphs[0]
        p.add_run(name + ': ').bold = True
        p.add_run(desc)
    sp()
    body(
        'Así mismo, conozco y acepto los términos del aviso de privacidad de la '
        'UAEMex, mismo que puede ser consultado en '
        'http://web.uaemex.mx/avisos/Aviso_Privacidad.pdf; en este mismo acto '
        'otorgo mi consentimiento, para que la Universidad Autónoma del Estado de '
        'México, haga públicos mis datos personales referentes a nombres, espacio '
        'académico, opiniones y/o conclusiones vertidas en el presente trabajo.'
    )

    # ── PÁGINA 4 – Firma ──
    doc.add_page_break()
    body(
        'En pos a la protección de datos personales de terceros, y en cumplimiento '
        'a la Ley de Protección de Datos Personales en Posesión de Sujetos '
        'Obligados, estoy de acuerdo para que la tesis de mi autoría no contenga '
        'documentos donde se visualicen datos personales sensibles que puedan '
        'afectar a terceros; tales documentos como voto aprobatorio, aceptación '
        'de tesis, dedicatorias, agradecimientos, mismos que, de no ocultarlos, '
        'serán visibles en el Repositorio Institucional de la Universidad Autónoma '
        'del Estado de México, haciéndome responsable de los mismos y sin previo '
        'permiso de los terceros.'
    )
    sp(); sp()
    _add_para(doc, 'Firmo de Conformidad y bajo protesta de decir verdad',
              align=WD_ALIGN_PARAGRAPH.CENTER)
    sp(); sp()
    _add_field(doc, 'Nombre y Firma', nombre)
    doc.add_paragraph('_' * 70)
    sp()
    _add_field(doc, 'No. De Cuenta', cuenta)
    doc.add_paragraph('_' * 70)
    sp(); sp()
    p_nota = doc.add_paragraph()
    r_n = p_nota.add_run(
        'NOTA: Ésta carta, toda vez que el autor registre los campos de llenado y '
        'las firmas correspondientes, debe digitalizarse y adjuntarse en el '
        'depósito del Repositorio Institucional de la Universidad Autónoma del '
        'Estado de México; misma que no será visible para consulta.'
    )
    r_n.bold = True; r_n.font.size = Pt(9); r_n.font.name = 'Courier New'
    sp()
    _add_para(doc,
              'Conozco y acepto los términos de privacidad de la\n'
              'Universidad Autónoma del Estado de México\n'
              'http://web.uaemex.mx/avisos/Aviso_Privacidad.pdf',
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── PÁGINA 5 – Hoja de datos del autor ──
    doc.add_page_break()
    right(f'Toluca, México a {fecha}')
    _add_heading(doc, 'Hoja de datos del autor')
    sp()
    _add_field(doc, 'Nombre:', nombre)
    _add_field(doc, 'Número de cuenta (en caso de aplicar):', cuenta)
    _add_field(doc, 'Grado académico:', director)
    _add_field(doc, 'Programa educativo de procedencia (aplica solo en tesis):', carrera)
    _add_field(doc, 'Institución donde labora:', '')
    _add_field(doc, 'Domicilio:', domicilio)
    _add_field(doc, 'Teléfono/Fax:', tel)
    _add_field(doc, 'Correo electrónico (preferentemente correo institucional):', correo)
    sp(); sp()
    p_firma = doc.add_paragraph('_' * 35)
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para(doc, 'Nombre y firma', align=WD_ALIGN_PARAGRAPH.CENTER)
    sp()
    p_n2 = doc.add_paragraph()
    r_n2 = p_n2.add_run(
        'Nota: para el caso de que sean más de un autor, se deberá imprimir esta '
        'última hoja de "datos del autor" en relación al número de autores.'
    )
    r_n2.font.size = Pt(9); r_n2.font.name = 'Courier New'
    sp()
    body('Esta información es recabada con fines administrativos')
    sp()
    _add_para(doc,
              'Conozco y acepto los términos de privacidad de la\n'
              'Universidad Autónoma del Estado de México\n'
              'http://web.uaemex.mx/avisos/Aviso_Privacidad.pdf',
              align=WD_ALIGN_PARAGRAPH.CENTER)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="carta_autorizacion_{cuenta}.docx"'
    return response


def lista_alumnos(request):
    alumnos = Alumno.objects.all()
    return render(request, 'alumnos/lista.html', {'alumnos': alumnos})
