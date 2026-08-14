"""Recibo de préstamo en Word: mismo patrón que
catalogo/constancia_donacion.py (documento construido desde cero con
python-docx, sin plantilla .docx)."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

VERDE = RGBColor(0x00, 0x66, 0x33)
GRIS = RGBColor(0x6B, 0x72, 0x80)


def generar_recibo_docx(prestamo):
    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('RECIBO DE PRÉSTAMO DE BIBLIOTECA')
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = VERDE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Biblioteca de Área de Medicina y Química — UAEMex')
    r.font.size = Pt(10)
    r.font.color.rgb = GRIS

    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(f'Folio: {prestamo.folio}')
    r.bold = True
    r.font.size = Pt(11)

    datos = [
        ('Alumno', prestamo.alumno_nombre),
        ('Matrícula', prestamo.matricula or '—'),
        ('Teléfono', prestamo.telefono or '—'),
        ('Carrera', prestamo.carrera or '—'),
        ('Fecha de salida', prestamo.fecha_salida.strftime('%d/%m/%Y')),
        ('Fecha de devolución', prestamo.fecha_devolucion.strftime('%d/%m/%Y')),
    ]
    for etiqueta, valor in datos:
        p = doc.add_paragraph()
        r = p.add_run(f'{etiqueta}: ')
        r.bold = True
        r.font.size = Pt(10)
        r = p.add_run(str(valor))
        r.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('Libros prestados')
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = VERDE

    items = list(prestamo.libros.select_related('ejemplar__registro'))
    tabla = doc.add_table(rows=1 + len(items), cols=2)
    tabla.style = 'Table Grid'
    hdr = tabla.rows[0].cells
    hdr[0].text = 'Código de barras'
    hdr[1].text = 'Título'
    for celda in hdr:
        celda.paragraphs[0].runs[0].bold = True
    for fila, item in zip(tabla.rows[1:], items):
        fila.cells[0].text = item.ejemplar.codigo_barras
        fila.cells[1].text = item.ejemplar.registro.titulo

    if prestamo.observaciones:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run('Observaciones')
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = VERDE
        doc.add_paragraph(prestamo.observaciones)

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('_' * 40)
    r.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Firma del alumno')
    r.font.size = Pt(9)
    r.font.color.rgb = GRIS

    return doc
