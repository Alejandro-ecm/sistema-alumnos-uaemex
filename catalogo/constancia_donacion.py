"""Generador de la constancia de recepción de donación de material
bibliográfico. Replica el diseño y contenido institucional que ya usaba
InventarioLibros (backend/pdf/constanciaWordGenerator.js): documento
construido desde cero con python-docx, sin plantilla .docx — mismo
patrón que alumnos/views.py (generar_carta)."""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

VERDE = RGBColor(0x00, 0x66, 0x33)
GRIS = RGBColor(0x6B, 0x72, 0x80)
GRIS_CLARO = RGBColor(0x9C, 0xA3, 0xAF)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

MESES = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
         'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']


def _fecha_larga(fecha):
    return f'{fecha.day} de {MESES[fecha.month - 1]} de {fecha.year}'


def _sombrear_celda(celda, color_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    celda._tc.get_or_add_tcPr().append(shd)


def _celda(tabla, fila, col, texto, negrita=False, color=None, fondo=None, tamano=10):
    celda = tabla.cell(fila, col)
    celda.text = ''
    run = celda.paragraphs[0].add_run(texto)
    run.bold = negrita
    run.font.size = Pt(tamano)
    if color:
        run.font.color.rgb = color
    if fondo:
        _sombrear_celda(celda, fondo)
    return celda


def generar_constancia_donacion_docx(constancia):
    items = list(
        constancia.libros
        .select_related('registro', 'registro__editorial')
        .prefetch_related('registro__autores')
    )
    titulos = len(items)
    volumenes = sum(item.cantidad for item in items)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f'No. {constancia.folio}   ')
    r.bold = True
    r.font.color.rgb = VERDE
    r.font.size = Pt(9)
    r = p.add_run(f'Toluca, Estado de México a {_fecha_larga(constancia.fecha)}')
    r.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('CONSTANCIA DE RECEPCIÓN DE DONACIÓN DE MATERIAL BIBLIOGRÁFICO')
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = VERDE

    p = doc.add_paragraph()
    r = p.add_run(constancia.persona_nombre.upper())
    r.bold = True
    r.font.size = Pt(11)

    if constancia.cargo.strip():
        p = doc.add_paragraph()
        r = p.add_run(constancia.cargo.strip())
        r.font.size = Pt(10)
        r.font.color.rgb = GRIS

    p = doc.add_paragraph()
    r = p.add_run('PRESENTE')
    r.bold = True
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(
        f'Agradecemos la donación del material documental de {volumenes} volumen(es) '
        f'correspondiente(s) a {titulos} título(s), mismo(s) que nos hizo llegar y que '
        f'enriquecerán el acervo de este Espacio Académico y sin duda serán de gran '
        f'beneficio a la comunidad universitaria.'
    )
    r.font.size = Pt(11)

    tabla = doc.add_table(rows=1 + len(items) + 1, cols=4)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = 'Table Grid'
    for col, texto in enumerate(['Título', 'Autor', 'Edición', 'Editorial']):
        _celda(tabla, 0, col, texto, negrita=True, color=BLANCO, fondo='006633')

    for i, item in enumerate(items, start=1):
        autores = ', '.join(a.nombre for a in item.registro.autores.all())
        editorial = item.registro.editorial.nombre if item.registro.editorial else ''
        fondo = None if i % 2 else 'f5f5f5'
        _celda(tabla, i, 0, item.registro.titulo, fondo=fondo)
        _celda(tabla, i, 1, autores, fondo=fondo)
        _celda(tabla, i, 2, item.registro.edicion, fondo=fondo)
        _celda(tabla, i, 3, editorial, fondo=fondo)

    fila_total = len(items) + 1
    celda_total = tabla.cell(fila_total, 0)
    celda_total.merge(tabla.cell(fila_total, 3))
    celda_total.text = ''
    p = celda_total.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f'Total de volúmenes recibidos: {volumenes}')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BLANCO
    _sombrear_celda(celda_total, '006633')

    p = doc.add_paragraph()
    r = p.add_run('Sin otro particular y agradeciendo su apoyo, le envío un cordial saludo.')
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ATENTAMENTE')
    r.bold = True
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('"Patria, Ciencia y Trabajo"')
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = GRIS

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('_' * 52)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('P. C.I.D. Marco Antonio Toledano')
    r.bold = True
    r.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Biblioteca de Área Dr. Rafael López Castañares')
    r.font.size = Pt(9)
    r.font.color.rgb = GRIS

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('BIBLIOTECA DE ÁREA DE MEDICINA Y QUÍMICA')
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = VERDE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'DOCUMENTO CONTROLADO EN EL SITIO WEB DEL SGC, QUE SE ENCUENTRA DISPONIBLE EXCLUSIVAMENTE '
        'PARA LA UNIVERSIDAD AUTÓNOMA DEL ESTADO DE MÉXICO. PROHIBIDA SU REPRODUCCIÓN TOTAL O PARCIAL.'
    )
    r.font.size = Pt(7)
    r.font.color.rgb = GRIS_CLARO

    return doc
